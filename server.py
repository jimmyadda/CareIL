import base64
from collections import defaultdict
from email import encoders
from email.mime.base import MIMEBase
import hmac
import io
import pathlib
from pydoc import text
import random
import secrets
import string
from bs4 import BeautifulSoup
import smtplib
import logging
from email.mime.multipart import MIMEMultipart
from email.mime.image import MIMEImage
from email.mime.text import MIMEText
import os
from dotenv import load_dotenv
from flask import Flask, abort, current_app, g, jsonify, send_file,flash,render_template,request,redirect, send_from_directory, session, url_for
import flask_login
import sqlite3
import datetime
import uuid
import hashlib
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt
from flask_restful import Resource, Api
from mail import get_Mail_settings, send_notification, update_Mail_setting
from package.decorators import admin_only
from package.patient import Patients, Patient
from package.doctor import Doctors, Doctor
from package.appointment import Appointments, Appointment,RequestAppointments,RequestAppointment
from package.common import Common
from package.User import User
from package.client import ClientUser
from package.medicalnote import Medicalnote,Medicalnotes
from flask_mail import Mail, Message
from create_account import create_account
from package.database import DatabaseManager
from package.google_calendar import (
    connection_status as google_calendar_connection_status,
    create_oauth_flow,
    disconnect as disconnect_google_calendar,
    is_configured as google_calendar_is_configured,
    save_connection as save_google_calendar_connection,
    sync_all_upcoming as sync_all_google_appointments,
)
from package.email_service import (
    careil_logo_attachment,
    encoded_attachment,
    resend_is_configured,
    resend_sender_address,
    send_resend_email,
)
from package.Myutils import render_ics
import json
from package.Auth2fa import store_verification_code,verify_code

load_dotenv()

app = Flask(__name__)

api = Api(app)
# Routes API
api.add_resource(Patients, '/patientapi')
api.add_resource(Patient, '/patientapi/<int:id>')
api.add_resource(Doctors, '/doctorapi')
api.add_resource(Doctor, '/doctorapi/<int:id>')
api.add_resource(Appointments, '/appointmentapi')
api.add_resource(Appointment, '/appointmentapi/<int:id>')
api.add_resource(RequestAppointments, '/appointmentrequestapi')
api.add_resource(RequestAppointment, '/appointmentrequestapi/<int:id>')
api.add_resource(Medicalnotes, '/medicalnoteapi')
api.add_resource(Medicalnote, '/medicalnoteapi/<int:id>')
api.add_resource(Common, '/common') 


with open('Translate.json',encoding="utf8") as Translate_file:
    Translate_data = json.load(Translate_file)

#Settings
with open('config.json') as config_file:
    config_data = json.load(config_file)
Globalsetting = config_data['Global'] 

 # Set up Flask-Mail
mail_settings = config_data['mail_settings']
app.config['MAIL_SERVER'] = mail_settings['MAIL_SERVER']  # Use your email provider's SMTP server
app.config['MAIL_PORT'] = mail_settings['MAIL_PORT']
app.config['MAIL_USE_TLS'] = mail_settings['MAIL_USE_TLS']
app.config['MAIL_USERNAME'] = os.environ.get('THERAPY_MAIL_USERNAME', mail_settings.get('MAIL_USERNAME', ''))
app.config['MAIL_PASSWORD'] = os.environ.get('THERAPY_MAIL_PASSWORD', mail_settings.get('MAIL_PASSWORD', ''))
mail = Mail(app) 



# Initialize DatabaseManager
db_manager = DatabaseManager()
# Ensure the default database is created at startup
db_manager.create_default_database()

SECRET_KEY = os.environ.get('THERAPY_SECRET_KEY', 'change-this-development-key')
app.config['SECRET_KEY'] = SECRET_KEY
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['SESSION_COOKIE_SECURE'] = os.environ.get('THERAPY_HTTPS', '').lower() in ('1', 'true', 'yes')
path = os.getcwd()

UPLOAD_FOLDER = os.path.join(path, 'uploads')
if not os.path.isdir(UPLOAD_FOLDER):
    os.mkdir(UPLOAD_FOLDER)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER


#Logs
handler = logging.FileHandler('LogFile.log') # creates handler for the log file
app.logger.addHandler(handler) # Add it to the built-in logger
app.logger.setLevel(logging.DEBUG)         # Set the log level to debug
logger = app.logger

EMAIL_LOGO_CID = "careil-logo"

def email_brand_header():
    return ("<div style='display:flex;align-items:center;gap:12px;margin:0 0 22px;"
            "padding:14px 18px;border-radius:14px;"
            "background:linear-gradient(135deg,#a3b18a 0%,#588157 100%);'>"
            "<img src='cid:careil-logo' alt='CareIL' "
            "style='display:block;width:64px;height:auto;'>"
            "<strong style='font:700 20px Arial,sans-serif;color:#102a1c;'>CareIL</strong>"
            "</div>")

def attach_email_logo(message):
    logo_path = os.path.join(os.path.dirname(__file__), 'static', 'img', 'therapy-hands-logo-email.png')
    with open(logo_path, 'rb') as logo_file:
        logo = MIMEImage(logo_file.read(), _subtype='png')
    logo.add_header('Content-ID', '<careil-logo>')
    logo.add_header('Content-Disposition', 'inline', filename='careil.png')
    message.attach(logo)
#Log in 
login_manager = flask_login.LoginManager()
login_manager.init_app(app)
#region Main App
base_db_path = "./databases/"

@login_manager.user_loader
def load_user(userid):  #or client patid
    user=None
    clientKey = session.get('client_key')
    if not clientKey:
        return None
    try:
        users = database_read(f"select * from accounts where userid='{userid}';",client_key=clientKey)
        client = database_read(f"select * from patient where pat_id='{userid}';",client_key=clientKey)
    except FileNotFoundError:
        # A deployment, database rename, or removed tenant can leave an old
        # browser cookie pointing at a database that no longer exists.
        session.pop('client_key', None)
        return None
    if len(users)==1 and bool(users[0].get('email_verified')):
        user = User(users[0]['userid'],users[0]['email'],users[0]['name'],users[0]['client_key'])
    if len(client)==1:
        user = ClientUser(client[0]['pat_id'],client[0]['pat_email'],client[0]['pat_first_name'],client[0]['client_key'])
    if user:
        session['client_key'] = user.client_key
        user.id = userid
        return user
    else:
        return None

def generate_client_key(user_id):
    # Use SHA-256 to generate a consistent hash
    return f"client_{hashlib.sha256(user_id.encode()).hexdigest()}" 

def ensure_single_therapist(client_key, user_data):
    """Create the one therapist profile required by appointment records."""
    conn = db_manager.connect_to_db(client_key=client_key)
    try:
        userid = user_data.get('userid', '')
        existing = conn.execute("SELECT doc_id FROM doctor WHERE userid=? LIMIT 1", (userid,)).fetchone()
        if existing:
            return existing['doc_id']

        # Older databases may have the single therapist row without the
        # account link. Attach it to the authenticated account instead of
        # creating a duplicate therapist.
        unlinked = conn.execute("SELECT doc_id FROM doctor LIMIT 1").fetchone()
        if unlinked:
            conn.execute("UPDATE doctor SET userid=? WHERE doc_id=?", (userid, unlinked['doc_id']))
            conn.commit()
            return unlinked['doc_id']

        full_name = (user_data.get('name') or user_data.get('userid') or 'Therapist').strip()
        name_parts = full_name.split(None, 1)
        first_name = name_parts[0]
        last_name = name_parts[1] if len(name_parts) > 1 else ''
        cursor = conn.execute(
            """INSERT INTO doctor
               (doc_first_name, doc_last_name, doc_email, doc_ph_no, doc_address, userid)
               VALUES (?, ?, ?, ?, ?, ?)""",
            (first_name, last_name, user_data.get('email', ''), '', '', userid)
        )
        conn.commit()
        return cursor.lastrowid
    finally:
        conn.close()
 
def database_write(sql,data=None):
    client_key = session['client_key']
    db_manager = DatabaseManager(client_key)
    conn = db_manager.connect_to_db(client_key)
    db = conn.cursor()

    row_affected = 0
    if data:
        row_affected = db.execute(sql, data).rowcount
    else:
        row_affected = db.execute(sql).rowcount
    #connection.commit()
    conn.commit()
    #db.close()
    #connection.close()

    return row_affected

def database_read(sql,data=None,client_key=None):
    if data and not client_key:
        if "userid" in data:
            client_key = generate_client_key(data['userid'])
            session['client_key'] =  client_key
    if client_key :#'client_key' in session:
         Curr_ClientKey = client_key  #Curr_ClientKey = session['client_key']
    else:
        Curr_ClientKey = Globalsetting['DEFAULT_CLIENT_KEY']
    print(Curr_ClientKey)
    conn = db_manager.connect_to_db(client_key=Curr_ClientKey)  # Connect to the client's database
    
    db = conn.cursor()   
    if data:
         db.execute(sql, data)
    else:
         db.execute(sql)
    records = db.fetchall()    
    rows = [dict(record) for record in records]
    #db.close()
    #connection.close()
    return rows

@app.teardown_appcontext
def close_db(exception=None):
    """
    Close the database connection at the end of each request.
    """
    db_manager.close_db_connection(exception)


#region index,login,register
@app.route("/")
def index_page():    
    default_user = db_manager.get_db_connection()
    default_user = db_manager.get_db_path()
    if 'default' in default_user:
        client_key = Globalsetting['DEFAULT_CLIENT_KEY']
    else:
        client_key = session['client_key']
    print(client_key)
    if flask_login.current_user.is_authenticated:
        logger.info(str(flask_login.current_user.get_dict()) + " Has Logged in")   
        user = flask_login.current_user.get_dict()
        #dump sql con
        apps = Appointments()
        appointments= apps.get()   
        return render_template('/index.html',Translate_data=Translate_data,user=user,appointments=appointments)
    else:
        return render_template("login.html",alert="", verification_step=0, email_step=0)

@app.route("/register", methods=['GET'])
def registration_page():
    return render_template('register.html',alert="", verification_step=0, email_step=0)

@app.route("/register", methods=['POST'])
def registration_request():
    form = dict(request.values)    
    folderid="0"
    if 'folderid' in request.values:
        folderid = request.values['folderid']
    id="1"
    if 'id' in request.values:
        id = request.values['id']
    reg_email = request.values['email']
    if reg_email:
        # Generate a unique client_key (e.g., UUID or hash)
        #client_key = f"client_{hash(form['userid'])}"
        client_key = generate_client_key(form['userid'])         
        form['client_key'] = client_key
        # Check if the client key/database already exists
        db_path = db_manager.get_db_path(client_key)
        if os.path.exists(db_path):
            return {"error": f"Client with key '{client_key}' already exists."}, 400
        # Create the client-specific database
        db_manager.create_client_database(client_key)  
        session['client_key'] =  client_key
        #create new connection
        checkconn= db_manager.connect_to_db(client_key)
        ok = create_account(form)
        ensure_single_therapist(client_key, form)
        session['formData'] = form
        print('ok:' ,ok)
        if ok == 1: 
            logger.info("New User Created: "+ form['name'])
            session['pending_verification_userid'] = form['userid']
            session['verification_step'] = 0
            session['email_step'] = 1               
            return render_template('register.html',alert="", verification_step=session['verification_step'], email_step=session['email_step'])          
        else:
            return redirect(f"/error") 
    else:
         return render_template('/register.html',alert = "Please insert valid email to register!", verification_step=0, email_step=0) 

@app.route("/login", methods=['GET'])
def login_page():
    return render_template('login.html',alert ="")

@app.route("/login", methods=['POST'])
def login_request():    
    form = dict(request.values)
    client_key = generate_client_key(form['userid'])    
    form['client_key'] = client_key          
    users = database_read("select * from accounts where userid=:userid",form,client_key=client_key)    
    formid = form['userid']
    if users :
        if len(users) == 1: #user name exist, password not checked
            salt = users[0]['salt']
            saved_key = users[0]['password']
            generated_key = hashlib.pbkdf2_hmac('sha256',form['password'].encode('utf-8'),salt.encode('utf-8'),10000).hex()            
            if saved_key == generated_key: #password match
                session['client_key'] = client_key
                if not bool(users[0].get('email_verified')):
                    flask_login.logout_user()
                    session['pending_verification_userid'] = formid
                    session['email_step'] = 1
                    session['verification_step'] = 0
                    return render_template(
                        'register.html',
                        alert="Please verify your email before signing in.",
                        verification_step=0,
                        email_step=1,
                    )
                user = load_user(formid)                 
                logger.info(f"Login successfull - '{formid}'  date: {str(datetime.datetime.now())}")
                # Store client_key in the session
                session['client_key'] = user.client_key            
                ensure_single_therapist(client_key, users[0])
                flask_login.login_user(user)
                return redirect('/')                        
            else: #password incorrect
                logger.info(f"Login Failed - '{formid}'  date: {str(datetime.datetime.now())}")
                return render_template('/login.html',alert = "Invalid user/password. please try again.") 
        else: #user name does not exist
            logger.info(f"Login Failed - '{formid}'  date: {str(datetime.datetime.now())}")
            return render_template('/login.html',alert = "Invalid user/password. please try again.")
        

@app.route('/enter_email', methods=['POST'])
def enter_email():
    client_key = session.get('client_key')
    pending_userid = session.get('pending_verification_userid')
    if not client_key or not pending_userid:
        flash("Your session expired. Please log in and try again.", "warning")
        return redirect(url_for('login_page'))

    email = request.form.get('email', '').strip()
    if not email:
        flash("Please enter a valid email address.", "danger")
        return redirect(url_for('registration_page'))
    session['email'] = email
    database_write(
        "UPDATE accounts SET email=? WHERE userid=?",
        (email, pending_userid),
    )
    
    # Generate a random verification code
    verification_code = ''.join(random.choices(string.ascii_uppercase + string.digits, k=6))

    # Store the verification code in the database with an expiration time (e.g., 10 minutes from now)
    expiration_time = datetime.datetime.now() + datetime.timedelta(minutes=5)
    store_verification_code(client_key, verification_code, expiration_time)

    # Send the verification code email
    if not send_verification_code(email, verification_code):
        session['email_step'] = 1
        session['verification_step'] = 0
        return render_template(
            'register.html',
            alert="The verification email could not be sent. Please try again.",
            verification_step=0,
            email_step=1,
        )

    # Store the code in the session to verify later
    session['verification_code'] = verification_code
    session['email_step'] = 0
    session['verification_step'] = 1
    return render_template(
        'register.html',
        alert="",
        email=email,
        verification_step=1,
        email_step=0,
    )

# Route for 2FA verification (User enters verification code)
@app.route('/verify', methods=['POST'])
def verify():
    client_key = session.get('client_key')
    pending_userid = session.get('pending_verification_userid')
    if not client_key or not pending_userid:
        flash("Your verification session expired. Please log in again.", "warning")
        return redirect(url_for('login_page'))

    entered_code = request.form.get('verification_code', '').strip()
    if not entered_code:
        return render_template(
            'register.html',
            alert="Please enter the verification code.",
            email=session.get('email', ''),
            verification_step=1,
            email_step=0,
        )

    # Verify the entered code
    print("entered_code",entered_code)
    if verify_code(client_key, entered_code):
        database_write(
            "UPDATE accounts SET email_verified=1 WHERE userid=?",
            (pending_userid,),
        )
        user = load_user(pending_userid)
        if not user:
            flash("The account could not be loaded after verification. Please sign in.", "warning")
            return redirect(url_for('login_page'))
        flask_login.login_user(user)
        session.pop('pending_verification_userid', None)
        session.pop('verification_code', None)
        session['verification_step'] = 0
        session['email_step'] = 0
        flash("Email verified successfully!", "success")
        return redirect('/')
    else:
        return render_template(
            'register.html',
            alert="Incorrect or expired code. Please try again.",
            email=session.get('email', ''),
            verification_step=1,
            email_step=0,
        )

@app.route("/logout")
@flask_login.login_required
def logout_page():
    session.pop('client_key', None)
    flask_login.logout_user()
    return redirect("/")

#endregion

#region General Route

@app.route("/error")
def error_page():
    return "there was an Error"

@app.route("/calendar")
@flask_login.login_required
def calendar_page():
    user = flask_login.current_user.get_dict()
    apps = Appointments()
    appointments = apps.get()
    duration = int(_availability_settings(user['client_key'])['APPOINTMENT_DURATION'])
    return render_template('calendar.html',user=user,appointments=appointments,appointment_duration=duration)

@app.route('/service-worker.js')
def service_worker():
    response = send_from_directory(app.static_folder, 'service-worker.js')
    response.headers['Cache-Control'] = 'no-cache'
    response.headers['Service-Worker-Allowed'] = '/'
    return response

@app.route('/send-mail', methods=['GET',"POST"])
def send_appointment(notification=''):
    user = flask_login.current_user.get_dict()    
    #Mail settings
    client_key = user['client_key']
    mail_settings = get_Mail_settings(client_key)
    app.config.update(mail_settings)
    mail = Mail(app)
    data = dict(request.values)
    my_lang = data['lang']
    print(my_lang)    
    pat_data =  database_read(f"select * from patient where pat_id= '{data['id']}';",client_key=client_key)
    doc_id = data['doc_id']
    doc_data =  database_read(f"select * from doctor where doc_id= '{doc_id}';",client_key=client_key)
    pat_id = pat_data[0]['pat_id']
    pat_email = pat_data[0]['pat_email']
    doc_fullname = doc_data[0]['doc_first_name']+" " + doc_data[0]['doc_last_name'] 
    pat_fullname = pat_data[0]['pat_first_name']+" "+pat_data[0]['pat_last_name']
    doc_address = doc_data[0]['doc_address']
    #gmail_url = add Appointment to calendar
    if my_lang=='HE':
        subject = "CareIL | פגישת טיפול עם: " + doc_fullname
    else:
        subject = "CareIL | Appointment with: " + doc_fullname
    sender_email = resend_sender_address() if resend_is_configured() else str(mail_settings['MAIL_USERNAME'])
    receiver_email = pat_email
    
    #Build Msg
    # Email content
    #Portal_url = request.host_url + f"/clients/client_login"
    base_url = request.host_url + "/portal"
    Portal_url = generate_patient_portal_url(base_url, pat_id, client_key)
    appointmentDuration = data['appointment_date']

    format = "%Y-%m-%dT%H:%M:%S.%fZ"
    Varformat = "%Y-%m-%d %H:%M:%S"

    date_obj = datetime.datetime.strptime(appointmentDuration, Varformat)
    duration = int(_availability_settings(client_key)['APPOINTMENT_DURATION'])
    time_change = datetime.timedelta(minutes=duration)
    appointmentEnd = date_obj + time_change 
    
    if my_lang=='HE':
        notification = 'נשמח לראותך '
        email_body = '''
                    <div id='App_mail' style="text-align: right;direction: rtl;" >
                    <b>נושא :</b> {Subject}<br>
                    <b>מטופל :</b> {Patient}<br>
                    <b>תאריך פגישה :</b> {Appointment Date}<br>
                    <br>
                    <b></b><br>
                    {note}
                    </div>
                    <!-- Button code -->
                    ''' 
    else:
        notification = 'We will be happy to see you. '
        email_body = '''
                    <div id='App_mail' style="text-align: right;direction: rtl;" >
                    <b>Subject :</b> {Subject}<br>
                    <b>Client:</b> {Patient}<br>
                    <b>Appointment date :</b> {Appointment Date}<br>
                    <br>
                    <b></b><br>
                    {note}
                    </div>
                    <!-- Button code -->
                    ''' 
    # Email data
    email_data = {
            'Subject': subject,    
            'Patient': pat_fullname,
            'Appointment Date': data['appointment_date'],
            'note': notification
            }          
    
    email_content = email_brand_header() + email_body.format(**email_data)
    email_content += f"<br><p><a href={Portal_url}>Log In to Portal</a></p>"

   

    # Create MIME message ICS File
    if my_lang=='HE':
        desc = u'פגישת טיפול'
    else:
        desc = u'Therapy session'        
    ics = render_ics(
            title=desc,
            description=desc,
            location= doc_address,
            start= date_obj,
            end= appointmentEnd,
            created=None,
            admin=sender_email,
            admin_mail=sender_email
        )
    if resend_is_configured():
        send_resend_email(
            receiver_email,
            subject,
            email_content,
            attachments=[
                careil_logo_attachment(os.path.dirname(__file__)),
                encoded_attachment(ics, "calendar.ics", "text/calendar"),
            ],
        )
        print('Email sent through Resend!')
        return redirect(f"/appointment")

    message = MIMEMultipart('related')    
    message['From'] = sender_email
    message['To'] = receiver_email    
    message['Subject'] = str(subject)
    message.attach(MIMEText(email_content,'html',_charset='utf-8')) 
    message.attach(MIMEText(email_content,'text/calendar',_charset='utf-8'))    
    attach_email_logo(message)
    #calendar
    
    attachment = MIMEBase('text', 'calendar; name=calendar.ics; method=REQUEST; charset=UTF-8')
    attachment.set_payload(ics.encode('utf-8'))
    encoders.encode_base64(attachment)
    attachment.add_header('Content-Disposition', 'attachment; filename=%s' % "calendar.ics")
    
    message.attach(attachment)
    # Connect to the SMTP server and send the email
    with smtplib.SMTP(mail_settings['MAIL_SERVER'], 587) as server:
        server.starttls()
        server.login(mail_settings['MAIL_USERNAME'], mail_settings['MAIL_PASSWORD'])        
        server.sendmail(sender_email, receiver_email, message.as_string().encode("UTF-8"))
        server.close()
    print('Email sent!')
    return redirect(f"/appointment")

def _google_calendar_redirect_uri():
    return os.environ.get('GOOGLE_REDIRECT_URI') or url_for('google_calendar_callback', _external=True)


@app.route('/admin/google-calendar')
@admin_only
def google_calendar_settings():
    user = flask_login.current_user.get_dict()
    configured = google_calendar_is_configured()
    connected = google_calendar_connection_status(user['client_key'], user['userid']) if configured else None
    return render_template(
        'google-calendar.html', user=user, configured=configured, connected=connected
    )


@app.route('/google-calendar/connect')
@admin_only
def google_calendar_connect():
    if not google_calendar_is_configured():
        return redirect(url_for('google_calendar_settings', error='Google credentials are not configured'))
    try:
        flow = create_oauth_flow(_google_calendar_redirect_uri())
        authorization_url, state = flow.authorization_url(
            access_type='offline', include_granted_scopes='true', prompt='consent'
        )
    except ImportError:
        return redirect(url_for('google_calendar_settings', error='Install the Google Calendar dependencies first'))
    session['google_oauth_state'] = state
    return redirect(authorization_url)


@app.route('/google-calendar/callback')
@admin_only
def google_calendar_callback():
    expected_state = session.pop('google_oauth_state', None)
    if not expected_state or request.args.get('state') != expected_state:
        abort(400, description='Invalid Google OAuth state')
    flow = create_oauth_flow(_google_calendar_redirect_uri(), state=expected_state)
    flow.fetch_token(authorization_response=request.url)
    user = flask_login.current_user.get_dict()
    save_google_calendar_connection(user['client_key'], user['userid'], flow.credentials)
    try:
        synced = sync_all_google_appointments(user['client_key'])
    except Exception:
        current_app.logger.exception('Initial Google Calendar sync failed')
        synced = 0
    return redirect(url_for('google_calendar_settings', connected='1', synced=synced))


@app.route('/google-calendar/sync', methods=['POST'])
@admin_only
def google_calendar_sync_now():
    user = flask_login.current_user.get_dict()
    return_to_calendar = request.form.get('next') == '/calendar'
    if not google_calendar_is_configured():
        return redirect(url_for(
            'google_calendar_settings',
            error='Google Calendar credentials are not configured yet.'
        ))
    if not google_calendar_connection_status(user['client_key'], user['userid']):
        return redirect(url_for(
            'google_calendar_settings',
            error='Connect Google Calendar before synchronizing appointments.'
        ))
    try:
        synced = sync_all_google_appointments(user['client_key'])
        if return_to_calendar:
            return redirect(url_for('calendar_page', synced=synced))
        return redirect(url_for('google_calendar_settings', synced=synced))
    except Exception:
        current_app.logger.exception('Manual Google Calendar sync failed')
        if return_to_calendar:
            return redirect(url_for('calendar_page', sync_error='Calendar sync failed. Reconnect Google Calendar.'))
        return redirect(url_for('google_calendar_settings', error='Calendar sync failed. Reconnect Google Calendar.'))


@app.route('/google-calendar/disconnect', methods=['POST'])
@admin_only
def google_calendar_disconnect():
    user = flask_login.current_user.get_dict()
    disconnect_google_calendar(user['client_key'], user['userid'])
    return redirect(url_for('google_calendar_settings', disconnected='1'))


@app.route('/admin/mail-settings', methods=['GET', 'POST'])
@admin_only
def mail_settings():

    client_key = session['client_key']
    db_manager = DatabaseManager(client_key)
    conn = db_manager.connect_to_db(client_key)

    if request.method == 'POST':
        # Update each setting from the form data
        for key in ['MAIL_SERVER', 'MAIL_PORT', 'MAIL_USE_TLS', 'MAIL_USERNAME', 'MAIL_PASSWORD']:
            if key in request.form and (key != 'MAIL_PASSWORD' or request.form[key]):
                update_Mail_setting(key, request.form[key])        
        flash("Mail settings updated successfully!", "success")
        return redirect(url_for('mail_settings'))

    # Fetch current settings to display in the form
    settings = get_Mail_settings(client_key)
    conn.close()
    return render_template('mail_settings.html', settings=settings)

@app.route('/SendNotification', methods=['POST'])
def Send_mail_Notification():
    data = dict(request.values)
    client_key = session['client_key']
    patien_id = database_read(f"select pat_id from patient WHERE pat_email ='{data['pat_email']}' order by pat_date desc LIMIT 1;",client_key=client_key)
    patien_data =  database_read(f"select * from patient WHERE pat_email ='{data['pat_email']}' order by pat_date desc LIMIT 1;",client_key=client_key)
    clinic_data =  database_read(f"select * from clinicinfo LIMIT 1;",client_key=client_key)
    data['client_key'] = client_key
    data['pat_id'] = patien_id[0]['pat_id']
    data['patien_data'] = patien_data
    data['clinic_data'] = clinic_data
    send_notification(data)
    return "ok"

@app.route('/admin/clinic-info', methods=['GET'])
@admin_only
def get_clinic_info():
    client_key = session['client_key']
    db_manager = DatabaseManager(client_key)
    conn = db_manager.connect_to_db(client_key)
    clinic_info = {}
    result = database_read("SELECT name, address, phone, email, website FROM clinicinfo LIMIT 1",client_key=client_key)
    if not result:
        return render_template('clinic-info.html', clinic=None, error="Clinic information not found"), 404
    return render_template('clinic-info.html', clinic=result[0])

@app.route('/admin/clinic-info', methods=['POST'])
@admin_only
def update_clinic_info():
    client_key = session['client_key']
    db_manager = DatabaseManager(client_key)
    conn = db_manager.connect_to_db(client_key)
    data = dict(request.values)
    cursor = conn.cursor()
  # Validate the input
    required_fields = ['name', 'address', 'phone', 'email']
    for field in required_fields:
        if field not in data or not data[field]:
             jsonify({"error": f"{field} is required"}), 400
    name =  data.get('name')
    address =  data.get('address')
    phone =  data.get('phone')
    email =  data.get('email')
    website =  data.get('website')
    # Construct the raw SQL update query
    clinic = database_read("SELECT * FROM clinicinfo LIMIT 1",client_key=client_key)
    if clinic:
        #Update
        query = f"UPDATE clinicinfo SET name ='{name}',address = '{address}',phone = '{phone}',email = '{email}',website = '{website}' WHERE id = 1"
        updateClinic = database_write(query,data)
        if updateClinic == 1 :
            flash("Clinic information updated successfully!", "success")
            return render_template('clinic-info.html', clinic=data)
    else:
      query = f"INSERT into  clinicinfo (name,address,phone,email,website) VALUES('{name}','{address}','{phone}','{email}','{website}');"
      ok = database_write(query,data)
      if ok == 1:
          flash("New Clinic information updated successfully!", "success")
          return render_template('clinic-info.html', clinic=data)

@app.route('/admin/adminPanel', methods=['GET'])
@admin_only
def admin_panel():
    return render_template('adminPanel.html')

def _availability_settings(client_key):
    defaults = {
        'AVAILABILITY_DAYS': '0,1,2,3,4',
        'AVAILABILITY_START': '08:00',
        'AVAILABILITY_END': '18:00',
        'APPOINTMENT_DURATION': '60',
    }
    conn = db_manager.connect_to_db(client_key=client_key)
    try:
        rows = conn.execute(
            "SELECT key, value FROM settings WHERE key IN (?, ?, ?, ?)",
            tuple(defaults.keys())
        ).fetchall()
        for row in rows:
            defaults[row['key']] = row['value']
        return defaults
    finally:
        conn.close()

@app.route('/api/availability', methods=['GET'])
def availability_api():
    client_key = session.get('client_key')
    if not client_key:
        return jsonify({'error': 'Clinic context is required.'}), 401
    settings = _availability_settings(client_key)
    return jsonify({
        'days': [int(day) for day in settings['AVAILABILITY_DAYS'].split(',') if day != ''],
        'start': settings['AVAILABILITY_START'],
        'end': settings['AVAILABILITY_END'],
        'duration': int(settings['APPOINTMENT_DURATION'])
    })

@app.route('/admin/availability', methods=['GET', 'POST'])
@admin_only
def availability_settings():
    client_key = session['client_key']
    if request.method == 'POST':
        selected_days = request.form.getlist('days')
        start = request.form.get('start', '08:00')
        end = request.form.get('end', '18:00')
        duration = request.form.get('duration', '60')
        if not selected_days or start >= end:
            flash('Choose at least one day and make sure the end time is after the start time.', 'danger')
        else:
            conn = db_manager.connect_to_db(client_key=client_key)
            try:
                values = {
                    'AVAILABILITY_DAYS': ','.join(selected_days),
                    'AVAILABILITY_START': start,
                    'AVAILABILITY_END': end,
                    'APPOINTMENT_DURATION': duration,
                }
                for key, value in values.items():
                    conn.execute(
                        "INSERT INTO settings (key, value) VALUES (?, ?) "
                        "ON CONFLICT(key) DO UPDATE SET value = excluded.value",
                        (key, value)
                    )
                conn.commit()
                flash('Availability saved.', 'success')
            finally:
                conn.close()
            return redirect(url_for('availability_settings'))
    settings = _availability_settings(client_key)
    return render_template(
        'availability.html',
        settings=settings,
        selected_days=settings['AVAILABILITY_DAYS'].split(',')
    )

#endregion

#region file add/delete/upload handeling

@app.route("/new-folder", methods=["POST"])
@flask_login.login_required
def create_new_folder():
    form = dict(request.values)
    id = str(uuid.uuid1())
    form['id'] = id
    sql = f"INSERT into folders (userid,id,name) VALUES (:userid,:id,:name);"
    ok = database_write(sql,form)
    if ok == 1:
       return "OK" 
    else:
       return "ERROR"

@app.route('/upload', methods=['POST'])
def upload(): 
        user = flask_login.current_user.get_dict() 
        data=  dict(request.values)
        clientKey = user['client_key'][:10]
        UPLOAD_FOLDER = os.path.join(path, 'uploads',clientKey)
        print(UPLOAD_FOLDER)
        if not os.path.isdir(UPLOAD_FOLDER):
            print("no dir")
            os.mkdir(UPLOAD_FOLDER)
        app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER 

        id = data['id']
        if 'file' not in request.files:
            flash('No file part')
            return redirect(request.url)

        file = request.files['file']

        if file.filename == '':
            flash('No file selected for uploading')
            return redirect(request.url)
        if file :                           
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            create = datetime.datetime.now().strftime("%Y-%m-%d")
            file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
            sql = f"INSERT into Patientfiles (pat_id,filename,filepath,createdate,userid) VALUES ('{id}','{filename}','{filepath}','{create}','{user['userid']}');"
            
            print("Upload_Sql",sql)
            ok = database_write(sql,data)
            if ok == 1:
               print('File successfully uploaded')
               return redirect(f"/patientform?id={id}")            
            else:
               return "ERROR"
        else:
            print('Allowed file types are txt, pdf, png, jpg, jpeg, gif')
            return redirect(request.url)

@app.route("/upload", methods=['GET'])
def upload_page():
    id = request.args.get('id')
    user = flask_login.current_user.get_dict()
    client_key = session['client_key'] 
    patientdata = database_read(f"select * from patient where pat_id= '{id}';",client_key = client_key)
    print("patientdata",patientdata)
    return render_template('upload.html',patientdata=patientdata)

@app.route('/delete_file', methods=['DELETE'])
@flask_login.login_required
def delete_file():
    user = flask_login.current_user.get_dict()
    data=  dict(request.values)
    clientKey = user['client_key'][:10]
    print("data",data)
    id = data['id']
    filename =  data['filename']
    sql = f"Delete from Patientfiles where pat_id ='{id}' and filename = '{filename}';"
    ok = database_write(sql,data)
    if ok == 1:
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        os.remove(filepath)
        print('File successfully Deleted')
        return render_template('patientform.html')                         
    else:
        return "ERROR" 

@app.route('/download_file/<path:filename>',methods=['GET',"POST"])
@flask_login.login_required
def download_file(filename):
    user = flask_login.current_user.get_dict()
    data=  dict(request.values)   
    client_key = session['client_key'] 
    myfile =  database_read(f"select * from Patientfiles where filename= '{filename}';" ,client_key = client_key)
    if myfile:
        str_path = myfile[0]['filepath']
        for root, dirs, files in os.walk(app.config['UPLOAD_FOLDER']):
            for name in files:            
                # As we need to get the provided python file, 
                # comparing here like this
                if name == filename:  
                    path = os.path.abspath(os.path.join(root, name))
                    uploads = os.path.join(app.root_path, app.config['UPLOAD_FOLDER'])
                    return send_from_directory(uploads, filename)
    else:
         return "Error"

@app.route('/delete_folder', methods=['DELETE'])
@flask_login.login_required
def delete_folder():    
    data=  dict(request.values)
    user = flask_login.current_user.get_dict() 
    id = data['folderid']
    sql = f"Delete from folders where id = '{id}';"
    print(sql)
    ok = database_write(sql,data)
    if ok == 1:
        print('project successfully Deleted')
        #Log
        logger.info(f"Project deleted - '{data['foldername']}' has been Sent by: {user['userid']} date: {str(datetime.datetime.now())}")
        return "OK"            
    else:
        return "ERROR" 
    
#endregion
   
#region Therapy Routes
@app.route("/doctor", methods=['GET'])
@app.route("/therapist", methods=['GET'])
@app.route("/myprofile", methods=['GET'])
@admin_only
def doctor_Page():
    id = request.args.get('id')
    user = flask_login.current_user.get_dict()
    return render_template('doctor.html',Translate_data=Translate_data,user=user)

@app.route("/patient", methods=['GET'])
@flask_login.login_required
def patient_Page():
    id = request.args.get('id')
    user = flask_login.current_user.get_dict()
    return render_template('patient.html',Translate_data=Translate_data,user=user)

@app.route("/appointment", methods=['GET'])
@admin_only
def appointment_Page():
        id = request.args.get('id')
        user = flask_login.current_user.get_dict()
        return render_template('appointment.html',Translate_data=Translate_data,user=user)

@app.route("/patientform", methods=['GET'])
@flask_login.login_required
def patient_folder_Load():
    id = request.args.get('id', type=int)
    if id is None:
        flash("Please select a patient first.", "warning")
        return redirect(url_for('patient_Page'))
    client_key = session['client_key']
    messages = database_read(f"select * from messages where pat_id= '{id}';",client_key=client_key)
    user = flask_login.current_user.get_dict()    
    patientdata = database_read(f"select * from patient where pat_id= '{id}';",client_key=client_key)
    tasksfiles = database_read(f"select * from Patientfiles where pat_id= '{id}';",client_key=client_key) #id = pat_id
    data = request.values
    if 'id' in request.values:
     id = request.values['id']
     #json.loads(data)
    if 'noteid' in request.values:
        noteid = request.values['noteid']
        med = Medicalnote()
        mednote = med.get(noteid)
    pat_id = id
    apps = Appointments()
    appointments = apps.getappointmentsbypatient(pat_id)    
    notes = Medicalnotes()
    pat_mednotes = notes.getnotebypatient(pat_id)    
    length = len(pat_mednotes)
    for i in range(length):
        # GET ONLY TEXT from DB
        test= pat_mednotes[i]['body']       
        data = json.loads(test)        
        content_html = data.get('content', '')
        soup = BeautifulSoup(content_html, 'html.parser')
        text = soup.get_text() 
        # Get text and split by <br> tag
        text_list = [tag.get_text() for tag in soup.find_all(['div', 'br'])] 
        if len(text_list)> 1:     
            text_with_separators = ','.join(text_list)       
            pat_mednotes[i]["text"]=text_with_separators 
        else:
            pat_mednotes[i]["text"]=text
    session['patientdata'] = patientdata
    print("patientdata",patientdata)
    if not patientdata:
        abort(404)
    return render_template('patientform.html',Translate_data=Translate_data,user=user,patient=patientdata[0],patientdata=patientdata,messages=messages,appointments=appointments,pat_mednotes=pat_mednotes,tasksfiles=tasksfiles, alert="")

@app.route('/patient/<int:pat_id>/summary.docx', methods=['GET'])
@flask_login.login_required
def export_patient_summary(pat_id):
    """Download a clean Word summary containing client details and session notes."""
    client_key = session['client_key']
    patients = database_read(
        "SELECT * FROM patient WHERE pat_id = ?",
        (pat_id,),
        client_key=client_key
    )
    if not patients:
        abort(404)

    notes = database_read(
        "SELECT create_date, body FROM medrecords WHERE pat_id = ? ORDER BY create_date DESC",
        (pat_id,),
        client_key=client_key
    )
    patient = patients[0]
    document = Document()
    styles = document.styles
    styles['Normal'].font.name = 'Aptos'
    styles['Normal'].font.size = Pt(11)

    document.add_heading('Therapy summary', 0)
    document.add_heading(
        f"{patient.get('pat_first_name', '')} {patient.get('pat_last_name', '')}".strip(),
        level=1
    )
    details = document.add_table(rows=0, cols=2)
    details.style = 'Light Shading Accent 1'
    for label, value in (
        ('Client ID', patient.get('pat_insurance_no', '')),
        ('Date of birth', patient.get('pat_dob', '')),
        ('Phone', patient.get('pat_ph_no', '')),
        ('Email', patient.get('pat_email', '')),
        ('Address', patient.get('pat_address', '')),
    ):
        row = details.add_row().cells
        row[0].text = label
        row[1].text = str(value or '')

    document.add_heading('Session notes', level=1)
    if not notes:
        document.add_paragraph('No session notes have been recorded.')
    for note in notes:
        document.add_heading(str(note.get('create_date') or 'Session'), level=2)
        body = note.get('body') or ''
        try:
            content_html = json.loads(body).get('content', '')
        except (TypeError, ValueError, AttributeError):
            content_html = body
        note_text = BeautifulSoup(str(content_html), 'html.parser').get_text('\n', strip=True)
        document.add_paragraph(note_text or 'No note content.')

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    safe_name = secure_filename(
        f"{patient.get('pat_first_name', 'client')}_{patient.get('pat_last_name', '')}_therapy_summary.docx"
    )
    return send_file(
        output,
        as_attachment=True,
        download_name=safe_name or 'therapy_summary.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route("/patientform", methods=["POST"])
@flask_login.login_required
def update_patien():
    user = flask_login.current_user.get_dict()
    client_key =  user['client_key']
    form = dict(request.values)
    id = form['pat_id']
    sql = "UPDATE patient SET pat_first_name =:pat_first_name, pat_last_name =:pat_last_name, pat_ph_no =:pat_ph_no, pat_address=:pat_address, pat_email =:pat_email, pat_insurance_no =:pat_insurance_no where pat_id =:pat_id"
    ok = database_write(sql,form)   
    if ok == 1:
        patientdata = database_read(f"select * from patient where pat_id= '{id}';",client_key=client_key)
        message = 'Success'
        return render_template('patientform.html',user=user,patient=patientdata,message=message)
    else:
       return "ERROR"

@app.route("/patientnotes" , methods=['GET'])
@flask_login.login_required
def patientnotes_page():
    user = flask_login.current_user.get_dict()
    data = request.values
    if 'id' in request.values:
     id = request.values['id']
     #json.loads(data)
    if 'noteid' in request.values:
        noteid = request.values['noteid']
        med = Medicalnote()
        mednote = med.get(noteid)
    pat_id = data['id']
    apps = Appointments()
    appointments = apps.get() 
    notes = Medicalnotes()
    pat_mednotes = notes.getnotebypatient(pat_id)
    #TEST GET ONLY TEXT
    test= pat_mednotes[0]['body']
    data = json.loads(test)
    content_html = data.get('content', '')
    soup = BeautifulSoup(content_html, 'html.parser')
    text = soup.get_text()
    print("text",text) 
    # Get text and split by <br> tag
    text_list = [tag.get_text() for tag in soup.find_all(['div'])]
    if len(text_list)>0:
        text_with_separators = ','.join(text_list)
        pat_mednotes[0]["text"]=text_with_separators
    else:
        pat_mednotes[0]["text"]=text
    return render_template('patientnotes.html',Translate_data=Translate_data,user=user,pat_mednotes=pat_mednotes)

@app.route("/medicalnote" , methods=['GET'])
@flask_login.login_required
@admin_only
def medicalnote_page():
    user = flask_login.current_user.get_dict()
    client_key = session['client_key']
    pat_id = request.args.get('id', type=int)
    if pat_id is None:
        flash("Please select a patient before creating a medical note.", "warning")
        return redirect(url_for('patient_Page'))
    patients = database_read("SELECT * FROM patient WHERE pat_id = ?", (pat_id,), client_key=client_key)
    if not patients:
        abort(404)
    apps = Appointments()
    appointments = apps.get() 
    texteditor = ""
    pat_mednotes=""
    if request.args.get('noteid'):
        noteid = request.args.get('noteid', type=int)
        med = Medicalnote()
        mednote = med.get(noteid)
        # notes = Medicalnotes()
        # pat_mednotes = notes.getnotebypatient(pat_id)   
        pat_mednotes = database_read("SELECT * FROM medrecords WHERE pat_id = ? AND rec_id = ?", (pat_id, noteid), client_key=client_key)
        if not pat_mednotes:
            abort(404)
        texteditor = json.loads(pat_mednotes[0]['body'] or '{}')
        session['textineditor'] = texteditor.get('content', '')
    else:
        session['textineditor'] = " "
    return render_template('medicalnote.html',Translate_data=Translate_data,user=user,patient=patients[0],appointments=appointments,pat_mednotes=pat_mednotes,texteditor=texteditor)

@app.route("/medicalnote" , methods=['POST'])
@flask_login.login_required
def updatemedicalnote():
    user = flask_login.current_user.get_dict()
    data = dict(request.values)
    client_key = session['client_key']
    id = data['pat_id']
    contentbdy = data['content']
    if 'noteid' in request.values:
        noteid = request.values['noteid']
        #update
        print("update:", noteid)
        now = datetime.datetime.now().strftime("%Y-%m-%d")
        sql = "UPDATE medrecords SET pat_id = ?, create_date = ?, body = ? WHERE rec_id = ?"
        ok = database_write(sql, (id, now, contentbdy, noteid))
        if ok == 1:
            return render_template('medicalnote.html',user=user,data=data)
        else:
            return "ERROR"
    else:
        #New   
        now = datetime.datetime.now().strftime("%Y-%m-%d")
        sql = "INSERT INTO medrecords (pat_id, create_date, body) VALUES (?, ?, ?)"
        ok = database_write(sql, (id, now, contentbdy))
        if ok == 1:
            return render_template('medicalnote.html',user=user,data=data)
        else:
            return "ERROR"

@app.route('/templates', methods=['GET'])
@flask_login.login_required
def get_templates_options():
    client_key = session['client_key']
    template = database_read(f"select rec_id,appointment_type from recordstamplates;",client_key=client_key)
    print("template",template)
    return jsonify({'templates': template})

@app.route('/Addappointment_type', methods=['POST'])
@flask_login.login_required
def create_type():
    data = request.json
    client_key = session['client_key']
    if not data.get('name'):
        return jsonify({"error": "Name is required"}), 400
    new_type = data['name']    
    try:
        sql = "INSERT INTO recordstamplates (appointment_type) VALUES (?)"
        ok = database_write(sql, (new_type,))
        if ok==1:
            return jsonify({"message": "appointment type created"}), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    return jsonify({"message": "appointment type created"}), 201

# Load template by appointment type
@app.route('/api/templates/<appointment_type>', methods=['GET'])
@flask_login.login_required
def get_template(appointment_type):
    client_key = session['client_key']
    template = database_read(
        "SELECT template FROM recordstamplates WHERE appointment_type = ?",
        (appointment_type,),
        client_key=client_key
    )
    if template:
        return jsonify({'template_text': template[0].get('template') or ''})
    return jsonify({'error': 'Template not found'}), 404

# Save or update a template
@app.route('/api/templates', methods=['POST'])
@flask_login.login_required
def save_template():
    data = request.get_json(silent=True) or {}
    client_key = session['client_key']
    appointment_type = (data.get('appointment_type') or '').strip()
    template_text = data.get('template_text')
    if not appointment_type or template_text is None:
        return jsonify({'error': 'Missing required fields'}), 400

    existing = database_read(
        "SELECT rec_id FROM recordstamplates WHERE appointment_type = ?",
        (appointment_type,),
        client_key=client_key
    )
    if existing:
        database_write(
            "UPDATE recordstamplates SET template = ? WHERE appointment_type = ?",
            (template_text, appointment_type)
        )
    else:
        database_write(
            "INSERT INTO recordstamplates (appointment_type, template) VALUES (?, ?)",
            (appointment_type, template_text)
        )
    return jsonify({'message': 'Template saved successfully'})
    
@app.route("/message" , methods=['GET'])
@flask_login.login_required
def message_page():
    user = flask_login.current_user.get_dict()
    data = request.values
    client_key = session['client_key']
    pat_id = data['patid']
    app_id = 0
    if 'app_id' in request.values:
        app_id =  data['app_id']
    if 'mark' in request.values:
        #Can Close messages only from portal
        pat_sign = session.get('client_key_signature')
        recid = data['rec_id']
        sql = f"update messages SET status=1 where rec_id = '{recid}';"
        ok = database_write(sql,data)
        if ok == 1:
            return redirect(url_for('get_portal'))
        else:
            return "ERROR"
        
    if 'rec_id' in request.values:
        recid = request.values['rec_id']   
        pat_messages = database_read(f"select * from messages where pat_id= '{pat_id}' and rec_id = '{recid}';",client_key=client_key)
    else:
        pat_messages = database_read(f"select * from messages where pat_id= '{pat_id}' ;",client_key=client_key)
        return render_template('message.html',user=user,pat_messages=pat_messages,pat_id=pat_id)

@app.route("/message" , methods=['POST'])
@flask_login.login_required
def updatemessages():
    user = flask_login.current_user.get_dict()
    data = dict(request.values)
    id = data['pat_id']
    app_id = 0
    if 'app_id' in request.values:
        app_id =  data['app_id']
    msg = data['msg']
    if 'rec_id' in request.values:
        recid = request.values['rec_id']
        #update
        now = datetime.datetime.now().strftime("%Y-%m-%d")
        sql = f"update messages SET pat_id= '{id}', create_date= '{now}' ,message = '{msg}',app_id='{app_id}' where rec_id = '{recid}';"
        ok = database_write(sql,data)
        if ok == 1:
            return render_template('medicalnote.html',user=user,data=data)
        else:
            return "ERROR"
    else:
        #New   
        now = datetime.datetime.now().strftime("%Y-%m-%d")
        sql = f"INSERT into messages (pat_id,create_date,message,app_id,status) VALUES  ('{id}','{now}','{msg}','{app_id}','0');"
        ok = database_write(sql,data)
        if ok == 1:
            return render_template('message.html',user=user,data=data)
        else:
            return "ERROR"
#endregion

#region Clients Routes

# Function to generate the signature
def generate_signature(pat_id, client_key, expires):
    data = f"{pat_id}:{client_key}:{expires}"
    return hmac.new(SECRET_KEY.encode(), data.encode(), hashlib.sha256).hexdigest()

def generate_patient_portal_url(base_url, pat_id, client_key):
    expires = int((datetime.datetime.now(datetime.timezone.utc) + datetime.timedelta(days=7)).timestamp())
    signature = generate_signature(pat_id, client_key, expires)
    return f"{base_url}?pat_id={pat_id}&client_key={client_key}&expires={expires}&signature={signature}"

def _portal_token_hash(token):
    return hashlib.sha256(token.encode('utf-8')).hexdigest()

def _create_portal_invitation(pat_id, client_key, created_by, lifetime_minutes=30):
    token = secrets.token_urlsafe(32)
    expires_at = (datetime.datetime.now(datetime.timezone.utc) +
                  datetime.timedelta(minutes=lifetime_minutes)).strftime('%Y-%m-%d %H:%M:%S')
    manager = DatabaseManager(client_key)
    conn = manager.connect_to_db(client_key)
    try:
        conn.execute(
            "INSERT INTO portal_invitations (pat_id, token_hash, expires_at, created_by) VALUES (?, ?, ?, ?)",
            (pat_id, _portal_token_hash(token), expires_at, created_by)
        )
        conn.commit()
    finally:
        conn.close()
    return token, expires_at

def _find_portal_invitation(token):
    """Find an opaque invitation without exposing the tenant key in its URL."""
    token_hash = _portal_token_hash(token)
    manager = DatabaseManager()
    if not os.path.isdir(manager.base_db_path):
        return None
    for filename in os.listdir(manager.base_db_path):
        if not filename.endswith('.db'):
            continue
        client_key = filename[:-3]
        try:
            conn = manager.connect_to_db(client_key)
            invitation = conn.execute(
                "SELECT * FROM portal_invitations WHERE token_hash = ? LIMIT 1", (token_hash,)
            ).fetchone()
            conn.close()
        except (sqlite3.Error, OSError):
            continue
        if invitation:
            invitation['client_key'] = client_key
            return invitation
    return None

@app.route('/api/patients/<int:pat_id>/portal-invitations', methods=['POST'])
@flask_login.login_required
@admin_only
def create_portal_invitation(pat_id):
    client_key = session['client_key']
    patient_rows = database_read(
        "SELECT pat_id, pat_first_name, pat_last_name, pat_email FROM patient WHERE pat_id = ?",
        (pat_id,), client_key=client_key
    )
    if not patient_rows:
        return jsonify({'error': 'Patient not found'}), 404
    patient = patient_rows[0]
    action = (request.get_json(silent=True) or {}).get('action', 'copy')
    if action not in ('copy', 'send'):
        return jsonify({'error': 'Invalid action'}), 400
    settings = None
    if action == 'send':
        if not patient.get('pat_email'):
            return jsonify({'error': 'This patient does not have an email address'}), 400
        settings = get_Mail_settings(client_key)
        if not resend_is_configured() and not str(settings.get('MAIL_USERNAME') or ''):
            return jsonify({'error': 'Mail settings are not configured'}), 400
    token, expires_at = _create_portal_invitation(
        pat_id, client_key, flask_login.current_user.get_id()
    )
    portal_url = url_for('portal_invitation_access', token=token, _external=True)
    if action == 'send':
        sender = str(settings.get('MAIL_USERNAME') or '')
        name = ' '.join(filter(None, [patient.get('pat_first_name'), patient.get('pat_last_name')]))
        subject = 'CareIL | Your secure client portal link'
        body = (email_brand_header() + f'<p>Hello {name},</p><p>Use the secure link below to open your client portal. '
                f'The link expires in 30 minutes and can be used once.</p>'
                f'<p><a href="{portal_url}">Open client portal</a></p>')
        if resend_is_configured():
            try:
                send_resend_email(
                    patient['pat_email'],
                    subject,
                    body,
                    attachments=[careil_logo_attachment(os.path.dirname(__file__))],
                )
                return jsonify({'url': portal_url, 'expires_at': expires_at, 'sent': True})
            except Exception:
                current_app.logger.exception('Failed to send portal invitation through Resend')
                database_write(
                    "UPDATE portal_invitations SET revoked_at = CURRENT_TIMESTAMP WHERE token_hash = ?",
                    (_portal_token_hash(token),)
                )
                return jsonify({'error': 'The portal email could not be sent. Check the Resend settings.'}), 502
        message = MIMEMultipart('related')
        message['From'] = sender
        message['To'] = patient['pat_email']
        message['Subject'] = subject
        message.attach(MIMEText(body, 'html', _charset='utf-8'))
        attach_email_logo(message)
        try:
            port = int(settings.get('MAIL_PORT') or 587)
            with smtplib.SMTP(settings['MAIL_SERVER'], port) as smtp:
                if str(settings.get('MAIL_USE_TLS', 'True')).lower() in ('1', 'true', 'yes'):
                    smtp.starttls()
                smtp.login(sender, settings['MAIL_PASSWORD'])
                smtp.sendmail(sender, patient['pat_email'], message.as_string().encode('utf-8'))
        except Exception:
            current_app.logger.exception('Failed to send portal invitation')
            database_write(
                "UPDATE portal_invitations SET revoked_at = CURRENT_TIMESTAMP WHERE token_hash = ?",
                (_portal_token_hash(token),)
            )
            return jsonify({'error': 'The portal email could not be sent. Check the mail settings.'}), 502
    return jsonify({'url': portal_url, 'expires_at': expires_at, 'sent': action == 'send'})

@app.route('/api/patients/<int:pat_id>/portal-invitations/revoke', methods=['POST'])
@flask_login.login_required
@admin_only
def revoke_portal_invitations(pat_id):
    updated = database_write(
        "UPDATE portal_invitations SET revoked_at = CURRENT_TIMESTAMP "
        "WHERE pat_id = ? AND used_at IS NULL AND revoked_at IS NULL",
        (pat_id,)
    )
    return jsonify({'revoked': updated})

@app.route('/portal/access')
def portal_invitation_access():
    token = request.args.get('token', '')
    invitation = _find_portal_invitation(token) if token else None
    if not invitation:
        return render_template('clientlogin.html', alert='This portal link is invalid.'), 403
    now = datetime.datetime.now(datetime.timezone.utc).strftime('%Y-%m-%d %H:%M:%S')
    if invitation.get('used_at') or invitation.get('revoked_at') or invitation['expires_at'] < now:
        return render_template('clientlogin.html', alert='This portal link has expired or has already been used.'), 403
    client_key = invitation['client_key']
    manager = DatabaseManager(client_key)
    conn = manager.connect_to_db(client_key)
    try:
        changed = conn.execute(
            "UPDATE portal_invitations SET used_at = CURRENT_TIMESTAMP "
            "WHERE invitation_id = ? AND used_at IS NULL AND revoked_at IS NULL AND expires_at >= ?",
            (invitation['invitation_id'], now)
        ).rowcount
        conn.commit()
    finally:
        conn.close()
    if changed != 1:
        return render_template('clientlogin.html', alert='This portal link is no longer available.'), 403
    session.clear()
    session['portal_authenticated'] = True
    session['portal_pat_id'] = invitation['pat_id']
    session['client_key'] = client_key
    return redirect(url_for('get_portal'))

# Function to validate the signature
def is_valid_signature(pat_id, client_key, expires, signature):
    try:
        if int(expires) < int(datetime.datetime.now(datetime.timezone.utc).timestamp()):
            return False
    except (TypeError, ValueError):
        return False
    expected_signature = generate_signature(pat_id, client_key, expires)
    return hmac.compare_digest(expected_signature, signature)   
     
@app.route("/clients/client_login", methods=['GET'])
def clientlogin_page():
    return render_template('/clientlogin.html',alert = "")      

@app.route("/clients/client_login", methods=['POST'])
def clientlogin_request():
    client_key = session["client_key"]
    form = dict(request.values)
    breakpoint()
    users = database_read(f"select * from patient where pat_email='{form['pat_mail']}' and pat_insurance_no='{form['pat_id']}';",client_key)
    print(users)
    print('users',users[0]['pat_id'])
    clientid= users[0]['pat_id']
    if len(users) >= 1: #user name exist, password not checked
        user = load_user(clientid)        
        flask_login.login_user(user)  
        return redirect(f"/portal?patid={clientid}") 
    else: #Invalid Email 
           return render_template('/clientlogin.html',alert = "Invalid Email/ID Number. please try again.")      
    
@app.route("/portal",methods=["GET"])
def get_portal():
    pat_id = session.get('portal_pat_id') if session.get('portal_authenticated') else None
    client_key = session.get('client_key') if pat_id else None
    expires = request.args.get("expires")
    signature = request.args.get("signature")
    appointment_dates= {}
    if not pat_id:
        pat_id = request.args.get("pat_id")
        client_key = request.args.get("client_key")
        if not pat_id or not client_key or not expires or not signature:
            return jsonify({"error": "Missing or expired portal access"}), 401
        if not is_valid_signature(pat_id, client_key, expires, signature):
            return jsonify({"error": "Invalid signature"}), 403
        session['portal_authenticated'] = True
        session['portal_pat_id'] = pat_id
    #user = flask_login.current_user.get_dict()
    session['client_key'] = client_key 
    if signature:
        session['client_key_signature'] = signature
    patientdata = database_read(f"select * from patient where pat_id= '{pat_id}';",None,client_key=client_key)
    patientmessages = database_read(f"select * from messages where status = 0 and pat_id= '{pat_id}';",None,client_key=client_key)
    lastappointment = database_read(f"SELECT *  FROM appointment where pat_id='{pat_id}' and appointment_date < DATETIME('now') order by appointment_date desc LIMIT 1;",None,client_key=client_key)
    nextappointment = database_read(f"SELECT *  FROM appointment where pat_id='{pat_id}' and appointment_date >= DATETIME('now') order by appointment_date  asc LIMIT 1;",None,client_key=client_key)
    if lastappointment:
        appointment_dates["lastappointment"] = lastappointment[0]["appointment_date"]
    if nextappointment:
        appointment_dates["nextappointment"] = nextappointment[0]["appointment_date"]
    apps = Appointments()
    appointments = apps.getappointmentsbypatient(pat_id)
    allappointments = apps.get()
    pending_count = len(patientmessages)
    session['patientdata'] = patientdata
    return render_template('portal.html',patientdata=patientdata,patientmessages=patientmessages,allappointments=allappointments,appointments=appointments,appointment_dates=appointment_dates,pending_count=pending_count,alert="")

@app.route('/portal/messages/<int:message_id>/read', methods=['POST'])
def portal_mark_message_read(message_id):
    pat_id = session.get('portal_pat_id') if session.get('portal_authenticated') else None
    client_key = session.get('client_key')
    if not pat_id or not client_key:
        return jsonify({'error': 'Portal session expired'}), 401
    updated = database_write(
        "UPDATE messages SET status = 1 WHERE rec_id = ? AND pat_id = ? AND status = 0",
        (message_id, pat_id)
    )
    if updated != 1:
        return jsonify({'error': 'Message not found'}), 404
    return jsonify({'message': 'Marked as read'})

@app.route('/portal/logout')
def portal_logout():
    session.clear()
    return redirect('/')

@app.route('/get-message/<int:message_id>', methods=['GET'])
def get_message(message_id):
    pat_id = session.get('portal_pat_id') if session.get('portal_authenticated') else None
    client_key = session.get('client_key')
    if not pat_id or not client_key:
        return jsonify({'error': 'Portal session expired'}), 401
    mymessages = database_read(
        "SELECT create_date, message FROM messages WHERE rec_id = ? AND pat_id = ?",
        (message_id, pat_id), client_key=client_key
    )
    if mymessages:
        return jsonify(mymessages)
    else:
        return jsonify({"error": "Message not found"}), 404
    

@app.route("/checkdate",methods=["POST"])
#@flask_login.login_required
def chekappointmentdate():
    data=  dict(request.values)
    #handel...
    print(data)
    #user = flask_login.current_user.get_dict() 
    id = data['pat_id']
    if session.get('portal_authenticated'):
        id = str(session.get('portal_pat_id'))
    datetocheck = data['appointmentdate']    
    client_key = session.get('client_key')
    try:
        requested_at = datetime.datetime.strptime(datetocheck, '%Y-%m-%d %H:%M:%S')
    except ValueError:
        return "ERROR", 400

    availability = _availability_settings(client_key)
    allowed_days = {int(day) for day in availability['AVAILABILITY_DAYS'].split(',') if day != ''}
    # JavaScript uses Sunday=0; Python uses Monday=0.
    requested_day = (requested_at.weekday() + 1) % 7
    start_time = datetime.datetime.strptime(availability['AVAILABILITY_START'], '%H:%M').time()
    end_time = datetime.datetime.strptime(availability['AVAILABILITY_END'], '%H:%M').time()
    if requested_day not in allowed_days or not (start_time <= requested_at.time() < end_time):
        return "ERROR"

    appoinmentindate = database_read(
        "SELECT app_id FROM appointment WHERE appointment_date = ? LIMIT 1",
        (datetocheck,),
        client_key=client_key
    )
    if len(appoinmentindate) >= 1:
       return "ERROR"             
    else:
        return "OK" 

@app.route("/postmsg",methods=["GET","POST"])
def postmsg():
    data=  dict(request.values)
    app_id = data['app_id']
    appointment = RequestAppointment()
    patapp = appointment.get(app_id)
    if patapp:
        pat_id=patapp[0]['pat_id']
        app_date = patapp[0]['appointment_date']
        msg="בקשתך לטיפול בתאריך : {app_date}  לא אושרה יש לבקש תאריך נוסף, יום נפלא".format(app_date=app_date)
        now = datetime.datetime.now().strftime("%Y-%m-%d")
        sql = f"INSERT into messages (pat_id,create_date,message,app_id) VALUES  ('{pat_id}','{now}','{msg}','{app_id}');"
        ok = database_write(sql,data)
        print('msg sent!',ok)
        return redirect(f"/appointment")  
    else:
        return "No Patient appointment"    
#endregion

def send_verification_code(email, verification_code):
    
    sender_email = resend_sender_address() if resend_is_configured() else app.config['MAIL_USERNAME']
    print(sender_email)
    """ Send the verification code to the user's email """
    msg = Message('CareIL | Your Verification Code', 
                  sender=sender_email, 
                  recipients=[email])
    msg.body = f'Your verification code is: {verification_code}. it will be valid for the next 5 Minutes.'
    msg.html = (email_brand_header() + '<p>Your verification code is:</p>'
                f'<p style="font-size:24px;font-weight:700;letter-spacing:4px;">{verification_code}</p>'
                '<p>It will be valid for the next 5 minutes.</p>')
    logo_path = os.path.join(os.path.dirname(__file__), 'static', 'img', 'therapy-hands-logo-email.png')
    with open(logo_path, 'rb') as logo_file:
        msg.attach(
            'careil.png', 'image/png', logo_file.read(),
            disposition='inline', headers=[['Content-ID', '<careil-logo>']]
        )
  
    
    try:
        if resend_is_configured():
            send_resend_email(
                email,
                'CareIL | Your Verification Code',
                msg.html,
                text=msg.body,
                attachments=[careil_logo_attachment(os.path.dirname(__file__))],
            )
        else:
            mail.send(msg)
        print(f"Verification email sent to {email}")
        return True
    except Exception as e:
        print(f"Failed to send email: {e}")
        return False


if __name__ == "__main__":
    app.run(
        host=Globalsetting.get("host", "127.0.0.1"),
        port=int(Globalsetting.get("port", 5000)),
        debug=os.environ.get("FLASK_DEBUG", "").lower() in ("1", "true", "yes")
    )
